import streamlit as st
import pandas as pd
import google.generativeai as genai
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

# API 키 보안 설정
# 페이지를 직접 실행할 경우 secrets.toml에서 키를 로드 시도
# 메인 앱에서 실행될 경우 st.session_state에 저장된 키를 사용
if 'user_api_key' in st.session_state and st.session_state.user_api_key:
    genai.configure(api_key=st.session_state.user_api_key)
else:
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception as e:
        st.error("Gemini API 키가 설정되지 않았습니다.")
        st.stop()

# 유효한 모델 이름으로 설정
model = genai.GenerativeModel('gemini-3.5-flash')


st.set_page_config(
    page_title="개별화교육계획 수립",
    page_icon="📄",
    layout="wide"
)

st.title("📄 AI 기반 개별화교육계획 수립 시스템")
st.markdown("---")

# --- IEP 생성 설정 ---
with st.container(border=True):
    st.header("📄 IEP 생성 설정")
    
    if 'curriculums' not in st.session_state:
        st.session_state.curriculums = ["기본교육과정"] 

    curriculums = st.multiselect(
        "1. 교육과정 선택 (복수 선택 가능)",
        ["기본교육과정", "공통교육과정"],
        default=st.session_state.curriculums
    )
    st.session_state.curriculums = curriculums

    subjects_by_curriculum = {
        "기본교육과정": ["국어", "수학", "생활영어", "진로와직업", "체육", "정보통신활용", "보건"],
        "공통교육과정": ["국어", "수학", "실과", "정보", "체육", "기술가정"]
    }
    
    curriculum_short_names = {
        "기본교육과정": "기본",
        "공통교육과정": "공통"
    }

    subject_to_curriculum = {}
    for cur, sub_list in subjects_by_curriculum.items():
        for sub in sub_list:
            if sub not in subject_to_curriculum:
                subject_to_curriculum[sub] = []
            subject_to_curriculum[sub].append(cur)

    available_subjects = set()
    if not curriculums:
        available_subjects.update(["국어", "수학"])
    else:
        for curriculum in curriculums:
            available_subjects.update(subjects_by_curriculum.get(curriculum, []))
    available_subjects = sorted(list(available_subjects))
    
    def format_subject(subject_name):
        if len(curriculums) > 1 and len(subject_to_curriculum.get(subject_name, [])) == 1:
            full_cur_name = subject_to_curriculum[subject_name][0]
            short_name = curriculum_short_names.get(full_cur_name, full_cur_name)
            return f"{subject_name} ({short_name})"
        return subject_name

    subject = st.selectbox(
        "2. 교과 선택", 
        options=available_subjects,
        format_func=format_subject,
        key="subject_selector"
    )
    st.session_state.subject = subject

    all_possible_grades = ["초등학교 1-2학년군", "초등학교 3-4학년군", "초등학교 5-6학년군", "중학교 1-3학년군"]
    available_grades = []
    if subject:
        grade_set = set()
        relevant_curriculums = [
            cur for cur, sub_list in subjects_by_curriculum.items() if subject in sub_list
        ]
        for curriculum in relevant_curriculums:
            for grade in all_possible_grades:
                file_path = f"data/{curriculum}/{subject}_{grade}.json"
                if os.path.exists(file_path):
                    grade_set.add(grade)
        available_grades = sorted(list(grade_set), key=all_possible_grades.index)

    grades = st.multiselect(
        "3. 학년군 선택",
        options=available_grades,
        default=[]
    )
    st.markdown("---")
    st.info("위 항목을 선택한 후, 아래 탭에서 단계를 진행하세요.")

tabs = st.tabs([
    "① 현행수준 진단", "② 현행수준 작성", "③ 교육목표 수립",
    "④ 교육내용 생성", "⑤ 교육 방법 선택", "⑥ 평가계획 수립", "⑦ 최종 IEP 생성"
])

# ---------------------------------------------------
# ① 현행수준 진단
# ---------------------------------------------------
with tabs[0]:
    if 'previous_grades' not in st.session_state:
        st.session_state.previous_grades = []
    if 'previous_subject' not in st.session_state:
        st.session_state.previous_subject = ""
    if 'previous_curriculums' not in st.session_state:
        st.session_state.previous_curriculums = []

    if (st.session_state.previous_grades != grades or
            st.session_state.previous_subject != subject or
            st.session_state.previous_curriculums != curriculums):
        
        keys_to_reset = ['evaluation', 'summary', 'goal_output', 'content_output', 'monthly_plan', 'selected_domains', 'evaluation_plan']
        for key in keys_to_reset:
            if key in st.session_state:
                del st.session_state[key]
        
        st.session_state.previous_grades = grades
        st.session_state.previous_subject = subject
        st.session_state.previous_curriculums = curriculums
        st.rerun()

    with st.container(border=True):
        st.header("① 현행수준 진단")
        
        domain_to_curriculum = {}
        criteria_by_domain = {}
        
        for curriculum in curriculums:
            if subject in subjects_by_curriculum.get(curriculum, []):
                for grade in grades:
                    file_path = f"data/{curriculum}/{subject}_{grade}.json"
                    if os.path.exists(file_path):
                        with open(file_path, "r", encoding="utf-8") as f:
                            try:
                                data = json.load(f)
                                for item in data:
                                    domain = item.get('영역', '기타')
                                    if domain not in domain_to_curriculum:
                                        domain_to_curriculum[domain] = set()
                                    domain_to_curriculum[domain].add(curriculum)
                                    
                                    if domain not in criteria_by_domain:
                                        criteria_by_domain[domain] = []
                                    item['출처'] = f"[{curriculum}] {grade}" 
                                    criteria_by_domain[domain].append(item)
                            except json.JSONDecodeError:
                                st.error(f"❌ JSON 파일 형식 오류: {file_path}")
                    else:
                        st.warning(f"⚠️ 성취기준 파일이 존재하지 않음: `{file_path}`")

        if not grades or not curriculums:
            st.info("IEP 생성 설정에서 진단할 교육과정과 학년군을 선택해주세요.")
        elif not domain_to_curriculum:
             st.warning("선택하신 조건에 해당하는 성취기준 파일이 없습니다. 설정을 확인해주세요.")
        else:
            st.subheader("1. 진단할 영역 선택")
            
            def format_domain(domain_name):
                if len(curriculums) > 1 and len(domain_to_curriculum.get(domain_name, set())) == 1:
                    full_cur_name = list(domain_to_curriculum[domain_name])[0]
                    short_name = curriculum_short_names.get(full_cur_name, full_cur_name)
                    return f"{domain_name} ({short_name})"
                return domain_name

            all_domains = sorted(list(domain_to_curriculum.keys()))
            if 'selected_domains' not in st.session_state:
                 st.session_state.selected_domains = all_domains
            
            selected_domains = st.multiselect(
                "이번 학기에 진단하고 계획을 수립할 영역을 선택하세요.",
                options=all_domains,
                format_func=format_domain,
                default=st.session_state.selected_domains
            )
            st.session_state.selected_domains = selected_domains
            
            st.markdown("---")
            st.subheader("2. 성취기준 기반 진단")

            if 'evaluation' not in st.session_state:
                st.session_state.evaluation = {}
            
            for domain in selected_domains:
                st.markdown(f"##### 🟦 {format_domain(domain)} 영역")
                items = criteria_by_domain.get(domain, [])
                for item in items:
                    key = f"[{item['출처']}] {item['id']}"
                    label_text = item['내용']
                    val = st.radio(label_text, ["예", "아니오", "관찰 필요"], key=key, horizontal=True)
                    st.session_state.evaluation[key] = {
                        "grade": item['출처'], "domain": domain, "id": item['id'],
                        "content": item['내용'], "value": val, "해설": item.get("해설", "")
                    }
    
    with st.container(border=True):
        st.subheader("🧐 '관찰 필요' 항목 진단 문항 생성")
        observation_needed = [v for v in st.session_state.get('evaluation', {}).values() if v.get('value') == "관찰 필요" and v.get('domain') in st.session_state.get('selected_domains', [])]
        if observation_needed:
            st.markdown("'관찰 필요'로 체크된 항목에 대해 학생의 현행 수준을 판단할 수 있는 객관적인 문항을 생성합니다.")
            if st.button("객관적 진단 문항 생성"):
                obs_text = "\n".join(f"- {v['content']}" for v in observation_needed)
                prompt_objective = f"""
                당신은 국가수준 학업성취도평가 문항을 출제하는 교육평가 전문가입니다.
                다음은 교사가 관찰만으로는 학생의 성취 여부를 판단하기 어려운 '관찰 필요' 항목들입니다.
                각 성취기준의 핵심 개념을 정확히 파악했는지 확인할 수 있는 **객관적인 평가 문항(선다형 또는 단답형)**을 각 항목당 1개씩 만들어주세요.
                **[성취기준 목록]**
                {obs_text}
                """
                with st.spinner('Gemini가 객관적 진단 문항을 생성하고 있습니다...'):
                    response = model.generate_content(prompt_objective)
                    obj_questions = response.text
                    st.success("📄 **생성된 객관적 진단 문항**")
                    st.markdown(obj_questions)
        else:
            st.info("현재 선택된 영역에서 '관찰 필요'로 체크된 항목이 없습니다.")

# ---------------------------------------------------
# ② 현행수준 작성
# ---------------------------------------------------
with tabs[1]:
    with st.container(border=True):
        st.header("② 현행수준 작성")
        if 'evaluation' in st.session_state and st.session_state.get('evaluation'):
            selected = {k: v for k, v in st.session_state.get('evaluation', {}).items() if v.get('value') == "예" and v.get('domain') in st.session_state.get('selected_domains', [])}
            if selected:
                st.markdown("✔️ **학생이 성취한 기준 요약:**")
                df = pd.DataFrame([{"학년군": v['grade'], "영역": v['domain'], "성취기준 ID": v['id'], "내용": v['content']} for v in selected.values()])
                st.dataframe(df, use_container_width=True, hide_index=True)
                st.markdown("---")
                st.markdown("🧠 **Gemini를 이용해 현행수준 요약문 생성**")
                if st.button("현행수준 문장 생성"):
                    input_text = "\n".join(f"- ({v['domain']} 영역) {v['content']}" for v in selected.values())
                    
                    prompt_template = f"""
                    당신은 특수교사를 돕는 IEP 작성 전문가입니다. 다음은 특수교육 대상학생의 {st.session_state.subject} 교과 성취기준 평가 결과 중 '예'로 체크된 항목입니다.
                    이를 바탕으로 학생의 강점을 보여주는 '현행학습수준'을 **하나의 자연스러운 종합 문단**으로 작성해 주세요.

                    **[학생이 성취한 기준 목록]**
                    {input_text}
                    
                    **[출력 규칙]**
                    - 각 영역(예: 읽기, 쓰기)의 강점들을 자연스럽게 연결하여 하나의 완성된 글로 작성하세요.
                    - **절대로 영역별로 목록을 나누거나 글머리 기호('-', '*')를 사용하지 마세요.**
                    - 학생의 강점을 나타내는 긍정적인 어조를 사용하세요.
                    - '~을 할 수 있으며, ~하는 능력을 보임.'과 같이 완전한 문장 형태로 자연스럽게 서술하세요.
                    """
                    with st.spinner('Gemini가 현행수준을 생성하고 있습니다...'):
                        response = model.generate_content(prompt_template)
                        summary = response.text.replace('*', '').replace('#', '').strip()
                        st.session_state.summary = summary
                
                if 'summary' in st.session_state:
                    st.success("📝 **Gemini 기반 현행학습수준 (아래 상자에서 수정 가능)**")
                    edited_summary = st.text_area(
                        "생성된 현행수준을 수정하거나 보완하세요.", 
                        value=st.session_state.summary, 
                        height=200,
                        key="summary_editor"
                    )
                    st.session_state.summary = edited_summary
            else:
                st.info("① 현행수준 진단 탭에서 학생이 성취한 기준('예')을 먼저 선택해주세요.")

# ---------------------------------------------------
# ③ 교육목표 수립
# ---------------------------------------------------
with tabs[2]:
    with st.container(border=True):
        st.header("③ 교육 목표 수립")
        if 'evaluation' in st.session_state and st.session_state.get('evaluation'):
            targets = [v for v in st.session_state.get('evaluation', {}).values() if v.get('value') in ["아니오", "관찰 필요"] and v.get('domain') in st.session_state.get('selected_domains', [])]
            if targets:
                st.markdown("✔️ **교육목표 수립 대상 (미도달 성취기준):**")
                df_targets = pd.DataFrame([{"학년군": v['grade'], "영역": v['domain'], "내용": v['content']} for v in targets])
                st.dataframe(df_targets, use_container_width=True, hide_index=True)
                st.markdown("---")
                st.markdown("🎯 **AI 기반 학기/월별 교육목표 자동 생성**")
                semester = st.radio("대상 학기 선택", ["1학기", "2학기"], horizontal=True, key="semester_radio")
                months_in_semester = {"1학기": ["3월", "4월", "5월", "6월", "7월"], "2학기": ["8월", "9월", "10월", "11월", "12월"]}
                selected_months = st.multiselect("목표를 생성할 월을 선택하세요", months_in_semester[semester], default=months_in_semester[semester])
                st.session_state.selected_months = selected_months
                if st.button("✏️ Gemini에게 교육목표 생성 요청"):
                    if not selected_months:
                        st.error("목표를 생성할 월을 1개 이상 선택해주세요.")
                    else:
                        criteria_text = "\n".join(f"- {v['id']} {v['content']}" for v in targets)
                        
                        prompt = f"""
                        당신은 IEP 교육목표를 작성하는 특수교육 전문가입니다.
                        
                        **[분석 자료]**
                        - 교과: {st.session_state.subject}, 대상 학기: {semester}, 목표 수립 월: {', '.join(selected_months)}
                        - 미도달 성취기준: {criteria_text}
                        
                        **[과업 지시]**
                        1. **학기 목표 생성**: 미도달 성취기준 전체를 아우르는 **{semester} 학기 목표**를 생성합니다.
                        2. **월별 목표 생성**: **{', '.join(selected_months)}** 각각에 해당하는 **월별 목표**를 구체적으로 생성합니다. 이때, 목표는 학생이 달성해야 할 '성취 상태'를 나타내도록 **'~할 수 있다', '~한다'** 와 같이 측정 가능한 **학생 중심**의 결과로 서술해 주세요.
                        
                        **[출력 형식 규칙]**
                        - 제목은 **'[1학기 학기 목표]', '[3월 목표]'와 같이 대괄호로 묶어서** 표시해주세요.
                        - **절대로 '#', '*'와 같은 다른 특수기호는 사용하지 마세요.**
                        - 각 월별 목표 다음 줄에는 '근거 성취기준:' 이라는 문구와 함께 관련 ID를 명시합니다.

                        **[출력 예시]**
                        [1학기 학기 목표]
                        일상생활 속 다양한 상황과 자료를 활용하여 자신의 생각과 느낌을 적절하게 표현하고, 타인과 바르고 고운 언어로 소통하며 즐겁게 국어 활동에 참여할 수 있다.

                        [3월 목표]
                        자신의 외모, 감정, 행동을 나타내는 간단한 단어와 짧은 문장을 사용하여 자신을 소개할 수 있다. 또한, 그림 자료를 통해 제시된 짧은 문장의 주요 내용을 파악할 수 있다.
                        근거 성취기준: 6국어01-02, 6국어02-03
                        """
                        with st.spinner('Gemini가 교육 목표를 생성하고 있습니다...'):
                            response = model.generate_content(prompt)
                            goal_output = response.text.replace('#### ', '').replace('### ', '')
                            st.session_state.goal_output = goal_output
                
                if 'goal_output' in st.session_state:
                    st.success("🧠 **Gemini 기반 학기/월별 목표 (아래 상자에서 수정 가능)**")
                    edited_goal = st.text_area(
                        "생성된 교육 목표를 수정하거나 보완하세요.",
                        value=st.session_state.goal_output,
                        height=400,
                        key="goal_editor"
                    )
                    st.session_state.goal_output = edited_goal
            else:
                st.info("선택하신 영역의 모든 성취기준을 성취했습니다.")

# ---------------------------------------------------
# ④ 교육내용 생성
# ---------------------------------------------------
with tabs[3]:
    with st.container(border=True):
        st.header("④ 교육내용 생성")
        if 'goal_output' in st.session_state:
            st.subheader("- 수정한 교육 목표 확인")
            st.markdown(st.session_state.goal_output)
            st.markdown("---")
            st.subheader("- 월별 교육내용 생성")
            if st.button("📚 Gemini에게 교육내용 생성 요청"):
                learning_goals_criteria = [v for v in st.session_state.get('evaluation', {}).values() if v.get('value') != "예" and v.get('domain') in st.session_state.get('selected_domains', [])]
                
                criteria_text_for_content = "\n".join(
                    f"- {v['id']} {v['content']}\n  (해설: {v.get('해설', '없음')})" for v in learning_goals_criteria
                )
                
                prompt_content = f"""
                당신은 학생 중심의 학습 활동을 설계하는 교육 전문가입니다. 아래 교육 목표를 달성하기 위해 학생이 직접 수행할 '주요 학습 활동' 목록을 생성해야 합니다.

                **[참고 자료]**
                1. **수립된 교육 목표:** {st.session_state.goal_output}
                2. **관련 성취기준 및 해설:** {criteria_text_for_content}

                **[과업 지시]**
                - 각 월별 목표를 달성하기 위한 **학생 중심의 주요 학습 활동을 3가지씩 제안**합니다.
                - 교사의 지도 내용이 아닌, 학생의 입장에서 수행하는 과제를 서술합니다.
                - **모든 활동 설명은 '~하기'와 같은 명사형으로 끝나야 합니다.** (예: '...답하는 활동을 합니다.' (X) -> '...답하기' (O))

                **[출력 형식 규칙]**
                - 각 월별 주요 학습 활동 섹션의 제목은 '### 3월 주요 학습 활동'과 같은 형식이어야 합니다.
                - 각 활동은 '**활동명:** 활동 설명' 형식으로 작성합니다.
                - **절대로 문장 앞에 `*`, `-`, `#` 와 같은 특수 기호를 사용하지 마세요.**
                - 각 활동은 반드시 줄을 바꿔서 작성합니다.

                **[출력 예시]**
                ### 3월 주요 학습 활동
                **주인공 되어보기:** 그림책이나 짧은 이야기 글을 읽고, 주인공이 되어 인터뷰 질문에 답하기
                **새로운 결말 상상하기:** 이야기의 결말을 자신만의 생각으로 새롭게 바꾸어 글이나 그림으로 표현하기
                **등장인물 관계도 그리기:** 이야기 속 등장인물들의 관계를 선과 간단한 설명으로 연결하여 한눈에 파악하기
                """
                with st.spinner('Gemini가 월별 교육내용을 생성하고 있습니다...'):
                    response = model.generate_content(prompt_content)
                    content_output = response.text
                    st.session_state.content_output = content_output
            
            if 'content_output' in st.session_state:
                st.success("🧠 **Gemini가 제안한 월별 지도 내용 및 방법 (아래 상자에서 수정 가능)**")
                edited_content = st.text_area(
                    "생성된 교육 내용을 수정하거나 보완하세요.",
                    value=st.session_state.content_output,
                    height=400,
                    key="content_editor"
                )
                st.session_state.content_output = edited_content
        else:
            st.info("③ 교육목표 수립 탭에서 먼저 교육목표를 생성하고 수정해주세요.")

# ---------------------------------------------------
# ⑤ 교육 방법 선택
# ---------------------------------------------------
with tabs[4]:
    with st.container(border=True):
        st.header("⑤ 교육 방법 선택")
        if 'goal_output' in st.session_state and 'content_output' in st.session_state:
            if 'monthly_plan' not in st.session_state:
                st.session_state.monthly_plan = {}

            def parse_monthly_data():
                goals_text = st.session_state.get('goal_output', '')
                contents_text = st.session_state.get('content_output', '')
                selected_months = st.session_state.get('selected_months', [])
                
                monthly_data = {month: {} for month in selected_months}

                goal_chunks = re.split(r'\[(\d{1,2}월) 목표\]', goals_text)[1:]
                for i in range(0, len(goal_chunks), 2):
                    month = goal_chunks[i]
                    if month in monthly_data:
                        goal_content = goal_chunks[i+1].strip()
                        if '근거 성취기준:' in goal_content:
                            goal_content = goal_content.split('근거 성취기준:')[0].strip()
                        monthly_data[month]['goal'] = goal_content
                
                content_chunks = re.split(r'### (\d{1,2}월) 주요 학습 활동', contents_text)[1:]
                for i in range(0, len(content_chunks), 2):
                    month = content_chunks[i]
                    if month in monthly_data:
                        monthly_data[month]['content'] = content_chunks[i+1].strip()

                for month in selected_months:
                    st.session_state.monthly_plan[month] = {
                        'goal': monthly_data.get(month, {}).get('goal', "파싱 실패: ③교육목표 탭을 확인해주세요."),
                        'content': monthly_data.get(month, {}).get('content', "파싱 실패: ④교육내용 탭을 확인해주세요."),
                        'methods': st.session_state.monthly_plan.get(month, {}).get('methods', []),
                        'other_method': st.session_state.monthly_plan.get(month, {}).get('other_method', "")
                    }
            
            parse_monthly_data()
            
            st.markdown("#### 월별 계획 및 교육 방법 선택")
            method_options = ["직접 교수법", "개념 학습법", "모델링 (시범)", "점진적 지원 감소", "협동학습 / 또래 교수", "기타 (직접 작성)"]

            for month, data in st.session_state.monthly_plan.items():
                with st.expander(f"**{month} 교육 계획 펼쳐보기**", expanded=True):
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        st.markdown("**월별 교육 목표**")
                        st.markdown(data.get('goal', '내용 없음'))
                        st.markdown("**주요 교육 내용**")
                        st.markdown(data.get('content', '내용 없음'))
                    with col2:
                        st.markdown("**교육 방법 선택**")
                        data['methods'] = st.multiselect(f"{month} 교육 방법", options=method_options, default=data['methods'], key=f"ms_{month}")
                        if "기타 (직접 작성)" in data['methods']:
                            data['other_method'] = st.text_area(f"{month} 기타 교육 방법", value=data['other_method'], key=f"ta_{month}")
            
            if st.button("월별 교육 방법 저장하기"):
                st.success("월별 교육 방법이 저장되었습니다!")
        else:
            st.info("③ 교육목표 수립 및 ④ 교육내용 생성 탭을 먼저 완료해주세요.")

# ---------------------------------------------------
# ⑥ 평가계획 수립
# ---------------------------------------------------
with tabs[5]:
    st.header("⑥ 평가계획 수립")
    if 'monthly_plan' not in st.session_state or not st.session_state.monthly_plan:
        st.info("⑤ 교육 방법 선택 탭에서 월별 계획을 먼저 수립하고 저장해주세요.")
    else:
        if 'evaluation_plan' not in st.session_state:
            st.session_state.evaluation_plan = {}
        
        for month in st.session_state.monthly_plan.keys():
            if month not in st.session_state.evaluation_plan:
                st.session_state.evaluation_plan[month] = {'methods': [], 'criteria': ''}

        EVAL_METHODS = [
            "관찰누가기록", "포트폴리오", "학습지/과제물 분석", 
            "질의응답", "발표", "프로젝트", "자기평가/동료평가"
        ]

        st.markdown("---")
        st.subheader("💡 AI 기반 평가계획 자동 생성")
        st.markdown("각 월별로 사용할 평가 방법을 먼저 선택한 후, '평가초점 생성' 버튼을 누르세요.")

        for month, plan_data in st.session_state.monthly_plan.items():
            with st.expander(f"**{month} 평가계획 수립**", expanded=True):
                
                current_eval_data = st.session_state.evaluation_plan[month]

                col1, col2 = st.columns([2, 1])

                with col1:
                    selected_methods = st.multiselect(
                        label=f"**{month} 평가 방법 선택**",
                        options=EVAL_METHODS,
                        default=current_eval_data.get('methods', []),
                        key=f"methods_{month}"
                    )
                    st.session_state.evaluation_plan[month]['methods'] = selected_methods

                with col2:
                    st.markdown("<br/>", unsafe_allow_html=True)
                    if st.button(f"**{month} 평가초점 생성**", key=f"btn_{month}", use_container_width=True):
                        if not selected_methods:
                            st.warning(f"{month} 평가 방법을 먼저 1개 이상 선택해주세요.")
                        else:
                            with st.spinner(f"Gemini가 {month} 평가초점을 생성하고 있습니다..."):
                                selected_methods_text = ", ".join(selected_methods)
                                
                                prompt_eval_plan = f"""
                                당신은 개별화교육계획(IEP) 전문가입니다.
                                아래는 학생의 월별 교육 목표와 내용이며, 이를 평가하기 위한 방법으로 '{selected_methods_text}'가 선택되었습니다.

                                - **월별 교육 목표**: {plan_data['goal']}
                                - **주요 교육 내용**: {plan_data['content']}
                                - **선택된 평가 방법**: {selected_methods_text}

                                **[과업 지시]**
                                선택된 평가 방법에 가장 적합한 **'평가 초점'**을 구체적인 질문 또는 확인 항목의 형태로 3~4가지 제안해 주세요.
                                
                                **[출력 규칙]**
                                - 마크다운 리스트(`- `) 형식으로 평가 초점만 간결하게 작성하세요.
                                - 각 항목은 학생의 성취 여부를 명확히 확인할 수 있는 내용이어야 합니다.
                                """
                                response = model.generate_content(prompt_eval_plan)
                                st.session_state.evaluation_plan[month]['criteria'] = response.text
                                st.success(f"{month} 평가초점 생성이 완료되었습니다!")

                if st.session_state.evaluation_plan[month].get('criteria'):
                    st.markdown("---")
                    st.markdown("##### 📝 **생성된 평가초점 (수정 가능)**")
                    edited_criteria = st.text_area(
                        label=f"생성된 {month} 평가초점을 수정하거나 보완하세요.",
                        value=st.session_state.evaluation_plan[month]['criteria'],
                        height=150,
                        key=f"criteria_{month}"
                    )
                    st.session_state.evaluation_plan[month]['criteria'] = edited_criteria

# ---------------------------------------------------
# ⑦ 최종 IEP 생성
# ---------------------------------------------------
with tabs[6]:
    st.header("⑦ 최종 IEP 미리보기 및 생성")

    st.subheader("1. 인적사항")
    col1_info, col2_info = st.columns(2)
    with col1_info:
        st.text_input("학생 이름", key="student_name")
    with col2_info:
        st.text_input("학년/반", key="student_class_info")

    st.subheader("2. 현재 학습 수행 수준")
    summary_text = st.session_state.get('summary', '② 현행수준 작성 탭에서 생성된 정보가 없습니다.')
    st.markdown(f"```\n{summary_text}\n```")

    st.subheader("3. 학기별 교육 계획")
    plan_data = []
    subject_name = st.session_state.get('subject', '')
    for month, data in st.session_state.get('monthly_plan', {}).items():
        methods_list = data.get('methods', [])
        other_method = data.get('other_method', '')
        if "기타 (직접 작성)" in methods_list and other_method:
            methods_list = [m if m != "기타 (직접 작성)" else other_method for m in methods_list]
        
        eval_data = st.session_state.get('evaluation_plan', {}).get(month, {})
        eval_methods = ", ".join(eval_data.get('methods', []))
        eval_criteria = eval_data.get('criteria', '').strip()
        eval_text = f"▪︎ 평가 방법: {eval_methods}\n▪︎ 평가 초점:\n{eval_criteria}"

        plan_data.append({
            "교과(영역)": f"{subject_name} ({month})",
            "장기 교육 목표 및 수립 근거": data.get('goal', ''),
            "교육 내용": data.get('content', ''),
            "교육 방법": ", ".join(methods_list),
            "평가 계획": eval_text
        })

    if plan_data:
        for month_plan in plan_data:
            with st.container(border=True):
                st.markdown(f"#### {month_plan['교과(영역)']}")
                st.markdown("**장기 교육 목표 및 수립 근거**")
                st.text(month_plan['장기 교육 목표 및 수립 근거'])
                st.markdown("**주요 교육 내용**")
                st.text(month_plan['교육 내용'])
                
                col_method, col_eval = st.columns(2)
                with col_method:
                    st.markdown("**교육 방법**")
                    st.text(month_plan['교육 방법'])
                with col_eval:
                    st.markdown("**평가 계획**")
                    st.text(month_plan['평가 계획'])
    else:
        st.info("③, ④, ⑤, ⑥ 탭을 진행하여 학기별 교육계획을 먼저 생성해주세요.")
    
    st.markdown("---")
    st.subheader("최종 IEP 문서(Word) 생성 및 다운로드")
    
    if st.button("📄 IEP 문서(Word) 생성 및 다운로드"):
        required_items = {
            'student_name': "학생 이름을 먼저 입력해주세요.", 'student_class_info': "학년/반을 먼저 입력해주세요.",
            'summary': "② 현행수준 작성 탭에서 내용을 생성하고 확인해주세요.", 'goal_output': "③ 교육목표 수립 탭에서 내용을 생성하고 확인해주세요.",
            'content_output': "④ 교육내용 생성 탭에서 내용을 생성하고 확인해주세요.", 'monthly_plan': "⑤ 교육 방법 선택 탭에서 내용을 확인하고 저장해주세요.",
            'evaluation_plan': "⑥ 평가계획 수립 탭에서 내용을 생성하고 확인해주세요."
        }
        all_ready = True
        for key, msg in required_items.items():
            if not st.session_state.get(key):
                st.error(msg); all_ready = False; break
        
        if all_ready:
            with st.spinner("IEP Word 문서를 생성 중입니다..."):
                document = Document()
                style = document.styles['Normal']; style.font.name = '맑은 고딕'; style.font.size = Pt(11)

                title = document.add_heading('개별화교육계획(IEP)', level=0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph()

                document.add_heading('1. 인적사항', level=1)
                table_info = document.add_table(rows=2, cols=4); table_info.style = 'Table Grid'
                table_info.cell(0, 0).text = '학생명'; table_info.cell(0, 1).text = st.session_state.get('student_name', '')
                table_info.cell(0, 2).text = '학년/반'; table_info.cell(0, 3).text = st.session_state.get('student_class_info', '')
                table_info.cell(1, 0).text = '교과'; table_info.cell(1, 1).text = st.session_state.get('subject', '')
                table_info.cell(1, 1).merge(table_info.cell(1, 3))
                for row in table_info.rows:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                    if len(row.cells) > 2: row.cells[2].paragraphs[0].runs[0].font.bold = True
                document.add_paragraph()

                document.add_heading('2. 현행학습수준', level=1)
                document.add_paragraph(st.session_state.get('summary', ''))
                document.add_paragraph()
                
                document.add_heading('3. 학기별 교육 계획', level=1)
                plan_table = document.add_table(rows=1, cols=5); plan_table.style = 'Table Grid'
                headers = ['교과(영역)', '교육 목표', '교육 내용', '교육 방법', '평가 계획']
                for i, header in enumerate(headers):
                    plan_table.rows[0].cells[i].text = header
                    plan_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                
                for month_plan in plan_data:
                    row_cells = plan_table.add_row().cells
                    row_cells[0].text = month_plan['교과(영역)']
                    row_cells[1].text = month_plan['교육 목표']
                    row_cells[2].text = month_plan['교육 내용']
                    row_cells[3].text = month_plan['교육 방법']
                    row_cells[4].text = month_plan['평가 계획']
                
                file_stream = io.BytesIO(); document.save(file_stream); file_stream.seek(0)
                st.success("✅ IEP 문서 생성이 완료되었습니다.")
                now_str = datetime.now().strftime("%Y%m%d")
                st.download_button(
                    label="📥 Word 파일(.docx) 다운로드",
                    data=file_stream,
                    file_name=f"IEP_{st.session_state.student_name}_{now_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# --- 저작권 표시 ---
st.markdown("---")
st.markdown("<p style='text-align: center; color: grey;'>Copyright © 2025 신하영(천안가온중학교), 성현준(청양고등학교). All Rights Reserved.</p>", unsafe_allow_html=True)