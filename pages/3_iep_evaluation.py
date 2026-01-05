import streamlit as st
import google.generativeai as genai
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- 🔄 세션 데이터 초기화 함수 ---
def reset_student_data():
    """새로운 학생 평가를 위해 모든 입력값과 세션 상태를 초기화함."""
    for key in list(st.session_state.keys()):
        if key not in ['user_api_key', 'GEMINI_API_KEY']:
            del st.session_state[key]
    st.rerun()

# --- 🔑 API 키 및 AI 모델 설정 ---
if 'user_api_key' in st.session_state and st.session_state.user_api_key:
    genai.configure(api_key=st.session_state.user_api_key)
else:
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception as e:
        st.error("Gemini API 키가 설정되지 않았음. 환경 변수나 사이드바 설정을 확인해야 함.")
        st.stop()

model = genai.GenerativeModel('gemini-2.0-flash')

st.set_page_config(
    page_title="AI 기반 개별화교육평가",
    page_icon="📝",
    layout="wide"
)

# --- ✨ 평가초점 생성 콜백 함수 (논리적 불일치 해결 및 서두 제거) ---
def generate_focus_callback(month, goal, content):
    if not goal or not content:
        st.error("평가초점을 생성하려면 먼저 해당 월의 교육 목표와 내용을 입력해야 함.")
        return

    # 프롬프트 설명: 수치(%) 배제, 행동 중심 서술, 서두/인사말 제거 강제
    prompt_focus = f"""
    당신은 특수교육 IEP 전문가임.
    교육 목표: {goal} / 교육 내용: {content}를 바탕으로 성취 수준을 관찰할 수 있는 '평가 초점' 5가지를 생성함.

    [절대 규칙 - 반드시 준수할 것]
    1. '20% 이상인가?', '3회 성공하는가?'와 같은 수치적 기준이나 성공 빈도는 절대 포함하지 마십시오. 
       도움의 수준(척도)과 매칭될 수 있도록 '특정 동작이나 기술의 수행 행위' 자체를 서술하십시오.
       (예: '슛 성공률이 20%인가?' -> '골 밑에서 골대를 향해 슛을 던지는 동작을 수행함')
    2. 항목만 바로 출력하십시오. "다음은 ~입니다"와 같은 서론, 인사말, 부연 설명은 절대 포함하지 마십시오.
    3. 모든 문장은 반드시 '~함' 또는 '~임'으로 끝나는 명사형 종결 어미를 사용하십시오.
    4. 각 항목을 줄바꿈으로 구분하여 리스트 형태로 출력하십시오.
    """
    with st.spinner(f"{month} 평가초점을 생성하는 중임..."):
        try:
            response = model.generate_content(prompt_focus)
            st.session_state[f"eval_focus_{month}"] = response.text.strip()
        except Exception as e:
            st.error(f"AI 생성 중 오류가 발생함: {e}")

# --- 🚀 UI 구성 시작 ---
st.title("📝 AI 기반 개별화교육평가")
st.markdown("---")
st.info("특수교육 IEP 평가의 전문성을 위해 모든 문장은 개조식(~함)으로 생성되며, 평가 초점은 척도와 일치하도록 행동 중심으로 설계됨.")

# 세션 상태 초기화
if 'evaluations_ai' not in st.session_state:
    st.session_state.evaluations_ai = {}
if 'semester_evaluation' not in st.session_state:
    st.session_state.semester_evaluation = {}

# --- ⭐ 성취도 척도 (6단계 전문 문체) ---
RATING_OPTIONS = [
    "도움 없이 스스로 과제를 완수함.",
    "시범을 보여주면 따라서 수행 가능함.",
    "한두 번의 언어적, 신체적 도움을 받으면 과제를 완수함.",
    "과제의 일부 단계를 도와주면 완수함.",
    "과제의 대부분 단계를 도와주어야 완수함.",
    "교사의 완전한 도움을 통해서만 과제 수행이 가능함."
]
RATING_SCORE_MAP = {opt: i+1 for i, opt in enumerate(RATING_OPTIONS)}

# --- 🚀 특이 상황별 전문 템플릿 ---
SPECIAL_CASE_TEMPLATES = {
    "외부 행사 등으로 인한 수업 시수 부족": "잦은 외부 활동 참여에 따른 수업 시수 부족으로 개별화교육계획에 수립된 내용을 계획대로 실시하지 못하였으며, 미진한 부분은 차기 교육과정에 반영하여 지속 지도하고자 함.",
    "치료 목적의 단축 수업(조퇴)": "건강 회복 및 외부 치료 지원을 위한 오전 단축 수업(등교 후 즉시 조퇴)이 지속됨에 따라, 실질적인 수업 참여 및 성취도 평가 근거가 미비함.",
    "잦은 지각 및 결석으로 인한 수업 미참여": "잦은 출결 변동(지각·결석)으로 인해 실질적인 수업 참여가 불규칙하여, 목표 달성 여부를 확인하기 위한 객관적인 평가 자료가 미비함."
}

with st.container(border=True):
    st.subheader("🗓️ 월별 교육 목표 입력 및 평가")
    semester = st.radio("평가 대상 학기 선택", ["1학기", "2학기"], horizontal=True, key="semester_radio_eval")
    months = {"1학기": ["3월", "4월", "5월", "6월", "7월"], "2학기": ["8월", "9월", "10월", "11월", "12월"]}[semester]
    
    for month in months:
        with st.container(border=True):
            st.subheader(f"✅ {month} 평가")
            
            # 운영 상황 선택
            status = st.selectbox(
                f"🚩 {month} 수업 운영 상황",
                ["정상 수업", "외부 행사 등으로 인한 수업 시수 부족", "치료 목적의 단축 수업(조퇴)", "잦은 지각 및 결석으로 인한 수업 미참여"],
                key=f"status_{month}"
            )
            
            goal_text = st.text_area(f"{month} 교육 목표", key=f"goal_{month}", height=80)
            instructional_text = st.text_area(f"{month} 교육 내용", key=f"instructional_{month}", height=100)
            
            # 정상 수업일 때만 평가 초점 및 척도 활성화
            if status == "정상 수업":
                col1, col2 = st.columns([4, 1])
                with col1:
                    eval_focus_text = st.text_area(f"{month} 평가초점 (행동 중심)", key=f"eval_focus_{month}", height=100)
                with col2:
                    st.write("") 
                    st.write("")
                    st.button(f"✨ 초점 생성", key=f"btn_gen_focus_{month}", on_click=generate_focus_callback, args=(month, goal_text, instructional_text))
                
                eval_focus_items = [item.strip() for item in eval_focus_text.split('\n') if item.strip()]

                if eval_focus_items:
                    st.markdown("#### 항목별 성취도 평가")
                    for i, item in enumerate(eval_focus_items):
                        st.markdown(f"**{i+1}. {item}**")
                        st.radio(
                            "성취도 선택", 
                            RATING_OPTIONS, 
                            key=f"rating_{month}_{i}", 
                            horizontal=True, 
                            label_visibility="collapsed"
                        )
                
                if st.button(f"🧠 {month} AI 종합 평가 생성", key=f"btn_ai_{month}"):
                    if not goal_text or not eval_focus_text:
                        st.error("목표와 평가초점을 입력해야 함.")
                    else:
                        full_eval_data = ""
                        for i, item in enumerate(eval_focus_items):
                            rating_label = st.session_state.get(f"rating_{month}_{i}", "평가되지 않음")
                            full_eval_data += f"- 평가 초점: {item} / 성취 수준: {rating_label}\n"

                        # 월별 평가 생성 시 '단순 나열' 방지를 위한 정교화된 프롬프트
                        prompt_eval = f"""
                        당신은 특수교육 전문가임. 제공된 자료로 학생 성취도를 전문적인 관찰 언어로 서술함.
                        [작성 규칙]
                        1. '초점+척도'를 단순히 합친 문장을 나열하지 마십시오.
                        2. 비슷한 수행 수준을 보인 항목들을 유기적으로 묶어서 하나의 문단으로 구성하십시오.
                        3. 강점과 보완점을 대조하는 연결어를 사용하여 문장의 흐름을 자연스럽게 만드십시오.
                        4. 모든 문장은 반드시 '~함', '~임', '~하였음'과 같은 명사형 종결 어미를 사용하십시오.
                        
                        목표: {goal_text}
                        관찰 데이터:
                        {full_eval_data}
                        """
                        with st.spinner(f"AI가 {month} 평가 문구를 생성 중임..."):
                            response = model.generate_content(prompt_eval)
                            st.session_state.evaluations_ai[month] = {
                                "goal": goal_text,
                                "instructional": instructional_text,
                                "evaluation": response.text
                            }
                        st.success(f"✔️ {month} 평가 문구 생성 완료!")

            # 특이 상황일 때 (시수 부족 등)
            else:
                st.warning(f"'{status}' 상황임. 아래 버튼을 클릭하여 전문 문구를 적용함.")
                if st.button(f"📋 {month} 특이사항 문구 적용", key=f"btn_special_{month}"):
                    st.session_state.evaluations_ai[month] = {
                        "goal": goal_text,
                        "instructional": instructional_text,
                        "evaluation": SPECIAL_CASE_TEMPLATES.get(status)
                    }
                    st.success(f"✔️ 특이사항 문구가 적용되었음.")

            # 최종 평가 결과 노출 및 편집
            if month in st.session_state.evaluations_ai:
                st.session_state.evaluations_ai[month]["evaluation"] = st.text_area(
                    f"{month} 최종 평가 문구 (편집 가능)",
                    value=st.session_state.evaluations_ai[month]["evaluation"],
                    key=f"ai_edit_{month}", height=150
                )

# ---------------- 🎓 학기 종합 평가 (요약 구조화 로직) ----------------
st.markdown("---")
st.subheader("🎓 학기 종합 평가")
if st.button("🧠 학기 종합 평가 생성", key="btn_semester_eval"):
    monthly_evals = {m: st.session_state.evaluations_ai[m] for m in months if m in st.session_state.evaluations_ai}
    if not monthly_evals:
        st.error("먼저 최소 한 달 이상의 평가를 생성해야 함.")
    else:
        # 월별 평가 내용을 하나의 텍스트로 합침
        full_semester_data = "\n\n".join([f"[{m} 평가 내용]\n{d['evaluation']}" for m, d in monthly_evals.items()])
        
        # 학기말 평가를 위한 강화된 프롬프트 (요약 및 구조화 요청)
        prompt_sem = f"""
        당신은 특수교육 전문가임. 제공된 학생의 월별 평가 내용을 분석하여 학기 전반의 성취를 종합 기술함.
        월별 내용을 각각 나열하지 말고, 전체 내용을 관통하는 공통적인 특성을 파악하여 아래의 4가지 항목으로 요약하여 작성하십시오.
        
        [작성 규칙]
        1. 모든 문장은 반드시 '~하였음.', '~할 수 있음.', '~가능함.', '~임.'과 같은 명사형 종결 어미로 작성함.
        2. 구조:
           - **강점 및 독립 수행 수준**: 한 학기 동안 학생이 스스로 수행 가능한 기술 및 두드러진 강점 요약.
           - **교사 지원을 통한 성취**: 시범이나 다양한 촉구(도움)를 통해 성공적으로 완수한 부분 요약.
           - **보완점 및 향후 지도 방향**: 여전히 어려움을 느끼는 부분과 이를 개선하기 위한 구체적인 지원 전략.
           - **최종 종합 의견**: 학생의 한 학기 전체 성취를 아우르는 전문적인 총평 한 문장.
        
        데이터:
        {full_semester_data}
        """
        with st.spinner("학기 종합 요약 평가 생성 중..."):
            response = model.generate_content(prompt_sem)
            st.session_state.semester_evaluation[semester] = response.text
        st.success("✔️ 학기 종합 요약 평가가 생성되었음!")

if st.session_state.semester_evaluation.get(semester):
    st.session_state.semester_evaluation[semester] = st.text_area(
        f"{semester} 종합 평가 편집",
        value=st.session_state.semester_evaluation[semester],
        key=f"semester_eval_editor", height=300
    )

# ---------------- 📥 결과 다운로드 및 초기화 ----------------
st.markdown("---")
st.subheader("📥 결과 다운로드 및 다음 학생 평가")
col_down, col_reset = st.columns(2)

with col_down:
    if st.button("📄 Word 파일 생성", key="btn_download_eval", use_container_width=True):
        with st.spinner("Word 파일을 생성하는 중임..."):
            document = Document()
            # 폰트 설정
            style = document.styles['Normal']
            style.font.name = '맑은 고딕'
            style.font.size = Pt(11)
            
            title = document.add_heading('개별화교육평가 결과 보고서', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}\n")
            
            for month in months:
                if month in st.session_state.evaluations_ai:
                    data = st.session_state.evaluations_ai[month]
                    document.add_heading(f"{month} 평가", level=2)
                    document.add_paragraph(f"▪︎ 교육 목표: {data['goal']}")
                    document.add_paragraph(f"▪︎ 주요 교육 내용:\n{data['instructional']}")
                    document.add_paragraph(f"▪︎ 종합 평가 결과:\n{data['evaluation']}\n")
            
            if semester in st.session_state.semester_evaluation:
                document.add_heading(f"{semester} 종합 요약 평가", level=1)
                document.add_paragraph(st.session_state.semester_evaluation[semester])

            # 메모리 내 저장 및 다운로드
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            
            st.download_button(
                label="📥 Word 파일 다운로드",
                data=file_stream,
                file_name=f"IEP_Evaluation_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

with col_reset:
    if st.button("🆕 새 학생 평가 시작 (데이터 초기화)", key="btn_main_reset", use_container_width=True, type="primary"):
        reset_student_data()
    st.caption("⚠️ 클릭 시 입력된 모든 데이터가 삭제되며 초기 상태로 돌아감.")

# --- 저작권 표시 ---
st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: grey;'>Copyright © 2026 신하영(천안가온중학교), "
    "성현준(청양고등학교). All Rights Reserved.</p>",
    unsafe_allow_html=True
)