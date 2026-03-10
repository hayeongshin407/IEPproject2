import io
from datetime import datetime

import google.generativeai as genai
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# =========================
# 페이지 설정
# =========================
st.set_page_config(
    page_title="개별화교육지원팀 협의회",
    page_icon="📝",
    layout="wide"
)


# =========================
# Gemini 설정
# =========================
def configure_gemini():
    api_key = None

    if "user_api_key" in st.session_state and st.session_state.user_api_key:
        api_key = st.session_state.user_api_key
    elif "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]

    if not api_key:
        st.warning(
            "Gemini API 키가 설정되지 않았습니다. "
            "메인 화면에서 API 키를 입력하거나, secrets.toml에 GEMINI_API_KEY를 설정해 주세요."
        )
        st.stop()

    genai.configure(api_key=api_key)
    return genai.GenerativeModel("gemini-2.0-flash")


model = configure_gemini()


# =========================
# 세션 상태 초기화
# =========================
def init_session_state():
    if "meeting_contents" not in st.session_state:
        st.session_state.meeting_contents = {
            "보호자 의견": "",
            "담임교사 의견": "",
            "특수교사 의견": "",
            "기타 의견": "",
            "의결 사항": ""
        }

    if "other_opinion_author" not in st.session_state:
        st.session_state.other_opinion_author = ""


init_session_state()


# =========================
# 공통 유틸
# =========================
def get_ai_refinement(prompt_text, content_type):
    if not prompt_text.strip():
        return ""

    if content_type == "의결 사항":
        prompt = f"""
당신은 개별화교육지원팀 회의록의 의결 사항을 전문가의 관점에서 명확하게 작성하는 역할을 합니다.
아래에 제시된 내용 요지를 바탕으로, 결정된 사항을 확정적으로 표현하는 개조식 문장으로 정리해 주세요.

[의결 사항 요지]
{prompt_text}

[출력 규칙]
- 마크다운 리스트(- ) 형식으로 각 항목을 정리하세요.
- 문장은 '~하기로 의결함', '~을 지원함'과 같이 확정적으로 마무리하세요.
- 오직 보완된 최종 문장만 출력하세요.
"""
    else:
        prompt = f"""
당신은 회의록 작성 전문가입니다. 아래에 제시된 의견 요지를 전문가의 어조로 다듬어 주세요.
내용을 간결하고 명확하게 정리하여 개조식 형태로 작성하고, 불필요한 내용은 제거해 주세요.

[회의 내용 요지]
{prompt_text}

[출력 규칙]
- 마크다운 리스트(- ) 형식을 사용하여 각 항목을 정리하세요.
- 각 항목의 문장은 '~이 필요함'으로 마무리하세요.
- 오직 보완된 최종 문장만 출력하세요.
"""

    try:
        response = model.generate_content(prompt)
        return response.text.strip() if response and response.text else prompt_text
    except Exception as e:
        st.error(f"AI 응답 생성 중 오류가 발생했습니다: {e}")
        return prompt_text


def queue_ai_refinement(input_key, content_type):
    current_text = st.session_state.get(input_key, "").strip()

    if not current_text:
        st.session_state[f"{input_key}_feedback"] = "내용을 먼저 입력해 주세요."
        return

    with st.spinner("AI가 내용을 보완하고 있습니다..."):
        refined_text = get_ai_refinement(current_text, content_type)

    st.session_state[f"{input_key}_pending"] = refined_text

    if refined_text.strip() == current_text.strip():
        st.session_state[f"{input_key}_feedback"] = (
            "보완 결과가 원문과 동일합니다. 입력 내용을 조금 더 구체적으로 작성해 보세요."
        )
    else:
        st.session_state[f"{input_key}_feedback"] = "AI 보완이 완료되었습니다."

    st.rerun()


def apply_pending_result(input_key, content_key):
    pending_key = f"{input_key}_pending"
    if pending_key in st.session_state:
        st.session_state[input_key] = st.session_state[pending_key]
        st.session_state.meeting_contents[content_key] = st.session_state[pending_key]
        del st.session_state[pending_key]


def render_feedback(input_key):
    feedback_key = f"{input_key}_feedback"
    if feedback_key in st.session_state:
        message = st.session_state[feedback_key]
        if message == "AI 보완이 완료되었습니다.":
            st.success(message)
        else:
            st.info(message)


def normalize_bullet_line(line: str) -> str:
    clean = line.strip()
    if clean.startswith("- "):
        clean = clean[2:].strip()
    return clean


def add_bullet_paragraphs(document, content: str):
    for line in content.split("\n"):
        clean_line = normalize_bullet_line(line)
        if clean_line:
            document.add_paragraph(clean_line, style="List Bullet")


def render_ai_refinement_section(
    title,
    expander_label,
    content_key,
    input_key,
    button_key,
    content_type,
    button_label="AI가 내용 보완하기"
):
    st.markdown(f"#### {title}")

    with st.expander(expander_label, expanded=True):
        if input_key not in st.session_state:
            st.session_state[input_key] = st.session_state.meeting_contents.get(content_key, "")

        apply_pending_result(input_key, content_key)

        st.text_area(
            "간단히 작성하면 AI가 보완해 줍니다.",
            key=input_key,
            height=150
        )

        st.session_state.meeting_contents[content_key] = st.session_state[input_key]

        if st.button(button_label, key=button_key):
            queue_ai_refinement(input_key, content_type)

        render_feedback(input_key)


# =========================
# 화면
# =========================
st.title("📝 개별화교육지원팀 협의회 회의록")
st.markdown("---")

with st.container(border=True):
    st.header("📋 회의 기본 정보")

    col1, col2 = st.columns(2)
    with col1:
        date_of_meeting = st.date_input("1. 일시 (날짜)", key="meeting_date")
    with col2:
        time_of_meeting = st.text_input("2. 일시 (시간)", key="meeting_time", value="14:00~15:00")

    location = st.text_input("3. 장소", key="meeting_location", value="특수교육지원실")

    meeting_type_options = ["서면 의견서 제출", "전화 상담", "대면 회의", "기타 (직접 작성)"]
    meeting_type = st.multiselect("4. 방식", meeting_type_options, default=["대면 회의"])

    other_method_text = ""
    if "기타 (직접 작성)" in meeting_type:
        other_method_text = st.text_input("기타 방식 내용을 입력하세요.", key="other_method_input")

    attendees = st.text_input(
        "5. 참석자 (쉼표로 구분하여 작성)",
        key="meeting_attendees",
        value="홍길동(담임), 김철수(특수교사), 이영희(보호자)"
    )

    st.markdown("---")
    st.header("📝 회의 내용")

    render_ai_refinement_section(
        title="보호자 의견 요지",
        expander_label="보호자 의견 작성",
        content_key="보호자 의견",
        input_key="parent_opinion_input",
        button_key="btn_parent_ai",
        content_type="의견"
    )

    render_ai_refinement_section(
        title="담임교사 의견 요지",
        expander_label="담임교사 의견 작성",
        content_key="담임교사 의견",
        input_key="teacher_opinion_input",
        button_key="btn_teacher_ai",
        content_type="의견"
    )

    render_ai_refinement_section(
        title="특수교사 의견 요지",
        expander_label="특수교사 의견 작성",
        content_key="특수교사 의견",
        input_key="special_teacher_opinion_input",
        button_key="btn_special_teacher_ai",
        content_type="의견"
    )

    st.markdown("#### 기타 의견 요지")
    with st.expander("기타 의견 작성", expanded=True):
        st.session_state.other_opinion_author = st.text_input("의견 제시자", key="other_author_input")

        if "other_opinion_input" not in st.session_state:
            st.session_state.other_opinion_input = st.session_state.meeting_contents["기타 의견"]

        apply_pending_result("other_opinion_input", "기타 의견")

        st.text_area(
            "간단히 작성하면 AI가 보완해 줍니다.",
            key="other_opinion_input",
            height=150
        )
        st.session_state.meeting_contents["기타 의견"] = st.session_state.other_opinion_input

        if st.button("AI가 내용 보완하기", key="btn_other_ai"):
            queue_ai_refinement("other_opinion_input", "의견")

        render_feedback("other_opinion_input")

    st.markdown("---")
    st.header("✅ 의결 사항")

    render_ai_refinement_section(
        title="의결 사항",
        expander_label="의결 사항 작성",
        content_key="의결 사항",
        input_key="resolution_input",
        button_key="btn_resolution_ai",
        content_type="의결 사항",
        button_label="AI가 의결 사항 보완하기"
    )

st.markdown("---")
st.subheader("📥 회의록 워드 파일 생성 및 다운로드")
st.markdown("입력된 내용을 바탕으로 워드 문서 파일을 생성합니다.")

if st.button("📄 워드 파일(.docx) 생성 및 다운로드"):
    if not st.session_state.meeting_contents["의결 사항"].strip():
        st.error("의결 사항을 먼저 작성해주세요.")
    else:
        with st.spinner("회의록 Word 문서를 생성 중입니다..."):
            document = Document()

            style = document.styles["Normal"]
            style.font.name = "맑은 고딕"
            style.font.size = Pt(11)

            title = document.add_heading("개별화교육지원팀 협의회 회의록", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph()

            # 1. 기본 정보
            p = document.add_paragraph("1. 협의회 기본 정보")
            p.runs[0].font.bold = True

            final_meeting_types = [m for m in meeting_type if m != "기타 (직접 작성)"]
            if other_method_text.strip():
                final_meeting_types.append(other_method_text.strip())

            table = document.add_table(rows=4, cols=2)
            table.style = "Table Grid"

            table.cell(0, 0).text = "일시"
            table.cell(0, 1).text = f"{date_of_meeting} {time_of_meeting}"

            table.cell(1, 0).text = "장소"
            table.cell(1, 1).text = location

            table.cell(2, 0).text = "방식"
            table.cell(2, 1).text = ", ".join(final_meeting_types)

            table.cell(3, 0).text = "참석자"
            table.cell(3, 1).text = attendees

            for row in table.rows:
                if row.cells[0].paragraphs and row.cells[0].paragraphs[0].runs:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True

            document.add_paragraph()

            # 2. 회의 내용
            p = document.add_paragraph("2. 회의 내용")
            p.runs[0].font.bold = True

            for section, content in st.session_state.meeting_contents.items():
                if section == "의결 사항":
                    continue
                if not content.strip():
                    continue

                section_title = (
                    f"- {st.session_state.other_opinion_author} 의견"
                    if section == "기타 의견" and st.session_state.other_opinion_author.strip()
                    else f"- {section}"
                )

                p = document.add_paragraph(section_title)
                p.runs[0].font.bold = True

                add_bullet_paragraphs(document, content)

            document.add_paragraph()

            # 3. 의결 사항
            p = document.add_paragraph("3. 의결 사항")
            p.runs[0].font.bold = True

            add_bullet_paragraphs(document, st.session_state.meeting_contents["의결 사항"])

            # 파일 저장
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)

            st.success("✅ 회의록 문서 생성이 완료되었습니다.")
            now_str = datetime.now().strftime("%Y%m%d")

            st.download_button(
                label="📥 Word 파일(.docx) 다운로드",
                data=file_stream,
                file_name=f"협의회_회의록_{now_str}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: grey;'>Copyright © 2025 신하영(천안가온중학교), "
    "성현준(청양고등학교). All Rights Reserved.</p>",
    unsafe_allow_html=True
)
